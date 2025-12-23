"""
Word文書マーカー抽出プログラム

【目的】
Word文書からマーカーで指定された範囲を抽出し、テキストファイルに出力する

【入出力】
- 入力: (プロジェクトルート)/data/input/*.docx
- 出力: (プロジェクトルート)/data/output/
  - 成功時: {元のファイル名}_{yyyymmddhhmmss}.txt
  - エラー時: {元のファイル名}_{yyyymmddhhmmss}_ERROR.txt
- ログ: (プロジェクトルート)/data/logs/process_{yyyymmddhhmmss}.log
- 複数ファイルを一括処理（処理開始時刻が全ファイル共通のタイムスタンプ）
- ~$で始まる一時ファイルは除外

【処理内容】
1. [PARENT]マーカーが出現したら、次の[PARENT]（または文末）までを出力
2. [SKIP]マーカーが出現したら、次のマーカー（[CHILD]/[SKIP]/[PARENT]）までを出力しない
3. [CHILD]マーカーが出現したら、[SKIP]状態を解除し出力を再開
4. マーカー自体も出力に含める

【読み込み対象】
本文（段落）、表、テキストボックスを文書内の出現順序通りに処理
※ 本文→表→本文のように混在していても、文書内の順序を保持

【表内マーカーの扱い】
表内に複数のマーカーが存在する場合、優先順位に従って1つだけ採用
- 優先順位: PARENT > SKIP > CHILD
- 理由: 表は2次元構造のため、マーカーの影響範囲（方向）を一意に判断できないため

【出力形式】
- 先頭以外の[PARENT]の前、すべての[CHILD]の前に空行を挿入
- 連続する空行は1行に統合
- テキストボックス・表セル内の改行を保持
- 表の各行は、セル内容を左から右へ ` | ` で区切って出力
- ファイル末尾は改行1つ

【エラー処理】
■ 部分的エラー（要素レベル）
  - 対象: 段落、表、テキストボックスの処理中のエラー
  - 動作: その要素のみをスキップし、同じファイル内の次の要素の処理を継続
  - ログ記録: エラー箇所、エラー内容、スタックトレースを詳細ログに記録

■ ファイル全体のエラー
  - 対象: ファイル読み込み失敗、ファイル破損などの致命的エラー
  - 動作: そのファイルをエラーファイル（_ERROR.txtサフィックス）として出力、
          次のファイルの処理を継続（全体処理は停止しない）
  - ログ記録: エラー内容をログに記録
"""

import traceback
import sys
from pathlib import Path
from datetime import datetime
from typing import Optional

import docx
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

# 終了コード（呼び出し元へ通知する契約）
EXIT_OK = 0  # 正常終了（全ファイル成功、警告相当なし）
EXIT_ERROR = 1  # 致命的エラー（環境不備などにより処理継続不可能）
EXIT_WARNING = 2  # 完走したが問題あり（警告、または _ERROR.txt 出力を伴うファイル単位失敗を含む）

# Word内のマーカー
MARKER_PARENT = "[PARENT]"
MARKER_CHILD = "[CHILD]"
MARKER_SKIP = "[SKIP]"

LINE_BREAK = "\n"  # 改行文字（出力用）

# ディレクトリ設定（プロジェクトルートからの相対パス）
# このスクリプトの位置: (project_root)/src/word/word_to_text.py
# プロジェクトルート = このスクリプトの2階層上
PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
INPUT_DIR = PROJECT_ROOT / "data" / "01_input"  # 入力ディレクトリ
OUTPUT_DIR = PROJECT_ROOT / "data" / "02_output"  # 出力ディレクトリ
LOG_DIR = PROJECT_ROOT / "logs"  # ログディレクトリ
FILE_PATTERN = "*.docx"  # 処理対象ファイルパターン

# グローバルログファイルハンドラ
log_file = None

def log(message, also_print=False):
    """ログメッセージをファイルに書き込む
    
    Args:
        message (str): ログメッセージ
        also_print (bool): コンソールにも出力するか
    """
    global log_file
    if log_file:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_file.write(f"[{timestamp}] {message}{LINE_BREAK}")
        log_file.flush()
    if also_print:
        print(message)

# XMLネームスペースを定義
# テキストボックス内のテキスト検索に必要
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'v': 'urn:schemas-microsoft-com:vml',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
}

# よく使うタグ名を定数化
TAG_W_R = f"{{{NSMAP['w']}}}r"
TAG_W_BR = f"{{{NSMAP['w']}}}br"

# --- ヘルパー関数 ---

had_warning = False      # 要素レベルのスキップ等
had_file_error = False   # ファイル単位の失敗（_ERROR.txt になるもの等）

def notify_warning(file_path: Optional[str], message: str):
    """
    要素レベルのワーニング（段落/表/テキストボックス等の部分的エラー）を通知する。

    - 処理は継続可能だが、当該要素はスキップされ出力結果が一部欠落する可能性がある。
    - 上位プロセスが機械的に検知できるよう、stderr に `WARNING:` で出力する。

    Args:
        file_path: 対象Wordファイルのパス（不明な場合は None/空文字列でも可）
        message: ワーニング内容（簡潔な要約）
    """ 
    global had_warning
    had_warning = True
    name = Path(file_path).name if file_path else "-"
    print(f"WARNING: {name}: {message}", file=sys.stderr)

def notify_file_error(file_path: Optional[str], message: str):
    """
    ファイル単位の失敗（当該ファイルが処理できず _ERROR.txt を出力する等）を通知する。

    - 当該ファイルは失敗扱いだが、全体処理は継続する。
    - 上位プロセスが機械的に検知できるよう、stderr に `ERROR:` で出力する。

    Args:
        file_path: 対象Wordファイルのパス
        message: エラー内容（簡潔な要約）
    """
    global had_file_error
    had_file_error = True
    name = Path(file_path).name if file_path else "-"
    print(f"ERROR: {name}: {message}", file=sys.stderr)

def notify_fatal(message: str):
    """
    致命的エラー（環境不備などで処理継続不能）を通知する。

    - 上位プロセスが機械的に検知できるよう、stderr に `FATAL:` で出力する。
    - 本関数は終了処理を行わない。呼び出し側が終了コード（例: 1）で終了することを想定する。

    Args:
        message: 致命的エラー内容（簡潔な要約）
    """
    print(f"FATAL: {message}", file=sys.stderr)

class ExtractionState:
    """
    マーカー抽出の状態を管理するクラス
    
    Attributes:
        in_parent (bool): [PARENT]セクション内かどうか
        in_skip (bool): [SKIP]セクション内かどうか
        found_parent_count (int): 検出した[PARENT]マーカーの数
    """
    def __init__(self):
        self.in_parent = False
        self.in_skip = False
        self.found_parent_count = 0
    
    def process_marker(self, marker_type):
        """マーカータイプに基づいて状態を更新
        
        Args:
            marker_type (str): MARKER_PARENT, MARKER_CHILD, MARKER_SKIP, または None
        """
        if marker_type == MARKER_PARENT:
            self.in_parent = True
            self.in_skip = False
            self.found_parent_count += 1
        elif marker_type == MARKER_CHILD:
            self.in_skip = False
        elif marker_type == MARKER_SKIP:
            self.in_skip = True

def check_marker_type(text):
    """テキスト中のマーカーを検出して、マーカータイプを返す
    
    Args:
        text (str): 検査対象のテキスト
    
    Returns:
        str or None: MARKER_PARENT, MARKER_CHILD, MARKER_SKIP のいずれか、またはマーカーがない場合は None
    """
    if MARKER_PARENT in text:
        return MARKER_PARENT
    elif MARKER_CHILD in text:
        return MARKER_CHILD
    elif MARKER_SKIP in text:
        return MARKER_SKIP
    return None

def get_combined_marker(*texts):
    """複数のテキストから最初に見つかったマーカーを返す
    
    段落本文とテキストボックス両方にマーカーがある場合、本文を優先
    
    Args:
        *texts (str): 検査対象のテキスト（可変長引数）
    
    Returns:
        str or None: 最初に見つかったマーカータイプ、またはマーカーがない場合は None
    """
    for text in texts:
        if text:
            marker = check_marker_type(text)
            if marker:
                return marker
    return None

def get_table_marker(table):
    """表内のマーカーを優先順位に従って取得
    
    表は2次元構造のため、複数マーカーがある場合は優先順位で1つだけ採用
    優先順位: PARENT > SKIP > CHILD
    
    Args:
        table (Table): 検査対象の表
    
    Returns:
        str or None: 採用されたマーカータイプ、またはマーカーがない場合は None
    """
    markers_found = set()
    for row in table.rows:
        for cell in row.cells:
            marker_type = check_marker_type(cell.text)
            if marker_type:
                markers_found.add(marker_type)
    
    # 優先順位に従って返す
    if MARKER_PARENT in markers_found:
        return MARKER_PARENT
    elif MARKER_SKIP in markers_found:
        return MARKER_SKIP
    elif MARKER_CHILD in markers_found:
        return MARKER_CHILD
    return None

def extract_drawingml_text(xml_element):
    """DrawingML形式のテキストボックスからテキストを抽出
    
    段落構造と改行を保持してテキストを抽出する
    XML解析エラーが発生した場合は空リストを返す
    
    Args:
        xml_element: 段落のXML要素
    
    Returns:
        list: 抽出されたテキストのリスト
    """
    texts = []
    
    try:
        # Wordの描画要素 (w:drawing) や図 (w:pict) の内部を探す
        for drawing in xml_element.findall('.//w:drawing', NSMAP) + xml_element.findall('.//w:pict', NSMAP):
            # 段落ごとに改行を保持
            paragraphs = drawing.findall('.//a:p', NSMAP)
            if paragraphs:
                para_texts = []
                for para in paragraphs:
                    para_text_parts = []
                    for text_tag in para.findall('.//a:t', NSMAP):
                        if text_tag.text:
                            para_text_parts.append(text_tag.text)
                    if para_text_parts:
                        para_texts.append(''.join(para_text_parts))
                if para_texts:
                    texts.append(LINE_BREAK.join(para_texts))
            else:
                # 段落構造がない場合は従来通り
                for text_tag in drawing.findall('.//a:t', NSMAP):
                    if text_tag.text:
                        texts.append(text_tag.text)
    except Exception as e:
        # Wordファイル内部のXML解析エラーを無視（破損したテキストボックスをスキップして処理継続）
        pass
    
    return texts

def extract_vml_text(xml_element):
    """VML形式のテキストボックスからテキストを抽出
    
    段落構造と改行タグ (w:br) を考慮してテキストを抽出する
    XML解析エラーが発生した場合は空リストを返す
    
    Args:
        xml_element: 段落のXML要素
    
    Returns:
        list: 抽出されたテキストのリスト
    """
    texts = []
    
    try:
        for textbox in xml_element.findall('.//v:textbox', NSMAP):
            # VMLの内部のWordprocessingMLの段落（w:p）を検索
            paragraphs = textbox.findall('.//w:p', NSMAP)
            
            # 段落が見つからない場合、txbxContentを探す
            if not paragraphs:
                for txbxContent in textbox.findall('.//w:txbxContent', NSMAP):
                    paragraphs = txbxContent.findall('w:p', NSMAP)
            
            if paragraphs:
                para_texts = []
                for para in paragraphs:
                    # 段落内のテキストランと改行を順番に処理
                    para_parts = []
                    for child in para:
                        # w:r (テキストラン) の場合
                        if child.tag == TAG_W_R:
                            # テキストランの中のテキスト
                            for text_run in child.findall('.//w:t', NSMAP):
                                if text_run.text:
                                    para_parts.append(text_run.text)
                            # 改行タグをチェック (直接の子要素のみ検索)
                            if child.findall('w:br', NSMAP):
                                para_parts.append(LINE_BREAK)
                    
                    if para_parts:
                        para_texts.append(''.join(para_parts))
                
                if para_texts:
                    texts.append(LINE_BREAK.join(para_texts))
            else:
                # 段落構造がない場合は従来通り
                for text_run in textbox.findall('.//w:t', NSMAP):
                    if text_run.text:
                        texts.append(text_run.text)
    except Exception as e:
        # Wordファイル内部のXML解析エラーを無視（破損したテキストボックスをスキップして処理継続）
        pass
    
    return texts

def extract_textbox_text(xml_element):
    """段落内のテキストボックスからテキストを抽出
    
    DrawingML形式とVML形式の両方に対応
    改行を保持してテキストを結合する
    
    Args:
        xml_element: 段落のXML要素
    
    Returns:
        str: 抽出されたテキスト（改行を含む）、またはテキストボックスがない場合は空文字列
    """
    texts = []
    texts.extend(extract_drawingml_text(xml_element))
    texts.extend(extract_vml_text(xml_element))
    return LINE_BREAK.join(texts) if texts else ''

def print_table(table):
    """表の全セル内容を行ごとにパイプ区切りで出力
    
    各行のセル内容を左から右へ ` | ` (スペース・パイプ・スペース) で区切る
    セル内の改行は保持される
    
    Args:
        table (Table): 出力対象の表
    """
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            # セル内の各段落を改行で結合（改行を保持）
            row_data.append(LINE_BREAK.join(p.text for p in cell.paragraphs))
        print(" | ".join(row_data))

def extract_marked_sections(file_name):
    """Word文書からマーカーで指定された範囲を抽出して出力
    
    [PARENT]セクション内かつ[SKIP]セクション外のコンテンツを出力する
    本文、表、テキストボックスを文書内の出現順序通りに処理する
    
    Args:
        file_name (str): 読み込むWord文書のファイル名
    
    Returns:
        int: 検出した[PARENT]マーカーの数
    """
    document = docx.Document(file_name)

    # 状態管理
    state = ExtractionState()
    last_was_empty = False  # 直前の出力が空行だったか
    is_first_output = True  # 最初の出力かどうか
    marker_just_output = False  # マーカーを出力した直後かどうか

    # 本文と表を順番通りに読み込む
    for element in document.element.body:
        
        if isinstance(element, CT_P):
            # 段落(本文)の場合
            try:
                paragraph = Paragraph(element, document)
                text = paragraph.text
                
                # --- テキストボックス抽出 ---
                textbox_text = extract_textbox_text(element)
                # ---------------------------
                
                # マーカーをチェック（本文とテキストボックスの両方）
                marker_type = get_combined_marker(text, textbox_text)
                
                # マーカーが見つかった場合、状態を更新
                if marker_type:
                    state.process_marker(marker_type)
                
                # 出力判定: [PARENT]セクション内 かつ [SKIP]セクション外
                if state.in_parent and not state.in_skip:
                    # マーカーの前に空行を挿入（先頭以外の[PARENT]、すべての[CHILD]）
                    # ただし、直前が空行の場合は挿入しない（連続空行防止）
                    if marker_type in [MARKER_PARENT, MARKER_CHILD] and not is_first_output and not last_was_empty:
                        print()
                        last_was_empty = True
                    
                    # 空行の連続を防ぐ（マーカー直後の空行も含む）
                    # ただし、テキストボックスがある場合は段落本文が空でも処理を継続
                    is_empty = (text == "" and not textbox_text)
                    if is_empty and (last_was_empty or marker_just_output):
                        # 連続する空行、またはマーカー直後の空行はスキップ
                        pass
                    else:
                        # 段落テキストを出力（ただし、テキストボックスのみの場合は段落本文の空行を出力しない）
                        if text or not textbox_text:
                            print(text)
                            is_first_output = False
                            
                            # last_was_emptyとmarker_just_outputを更新
                            if marker_type in [MARKER_PARENT, MARKER_CHILD]:
                                # マーカーを出力した場合
                                marker_just_output = True
                                last_was_empty = False  # マーカーは空行ではない
                            elif text == "":
                                # 空行を出力した場合
                                last_was_empty = True
                                marker_just_output = False
                            else:
                                # 通常テキストを出力した場合
                                last_was_empty = False
                                marker_just_output = False
                        
                    # テキストボックスがあれば追加で出力
                    if textbox_text:
                        print(textbox_text)
                        last_was_empty = False
                        is_first_output = False
                        marker_just_output = False
            
            except Exception as e:
                # 段落処理エラー時はその段落をスキップして処理継続
                error_detail = traceback.format_exc()
                log(f"  エラー: 段落処理で例外が発生 - {str(e)}")
                log(f"  スタックトレース:{LINE_BREAK}{error_detail}")
                # エラー発生段落の内容をログ出力
                try:
                    log(f"    エラー発生段落テキスト: {text}")
                    log(f"    エラー発生テキストボックス: {textbox_text}")
                except Exception as log_e:
                    log(f"    エラー発生段落内容のログ出力にも失敗: {log_e}")

                # 要素レベル警告として上位へ通知
                notify_warning(file_name, f"段落処理をスキップ: {e}")

                pass
        
        elif isinstance(element, CT_Tbl):
            # 表の場合
            try:
                table = Table(element, document)
                
                # 表のセル内でマーカーをチェック（優先順位付き）
                table_marker = get_table_marker(table)
                
                # マーカーが見つかった場合、状態を更新
                if table_marker:
                    state.process_marker(table_marker)
                
                # 出力判定: [PARENT]セクション内 かつ [SKIP]セクション外
                if state.in_parent and not state.in_skip:
                    # マーカーの前に空行を挿入（先頭以外の[PARENT]、すべての[CHILD]）
                    # ただし、直前が空行の場合は挿入しない（連続空行防止）
                    if table_marker in [MARKER_PARENT, MARKER_CHILD] and not is_first_output and not last_was_empty:
                        print()
                        last_was_empty = True
                    
                    print_table(table)
                    last_was_empty = False
                    is_first_output = False
                    
                    # 表にマーカーがある場合はフラグを立てる、そうでなければリセット
                    if table_marker in [MARKER_PARENT, MARKER_CHILD]:
                        marker_just_output = True
                    else:
                        marker_just_output = False
            
            except Exception as e:
                # 表処理エラー時はその表をスキップして処理継続
                error_detail = traceback.format_exc()
                log(f"  エラー: 表処理で例外が発生 - {str(e)}")
                log(f"  スタックトレース:{LINE_BREAK}{error_detail}")
                # エラー発生表の内容をログ出力
                try:
                    for row_idx, row in enumerate(table.rows):
                        row_texts = [LINE_BREAK.join(p.text for p in cell.paragraphs) for cell in row.cells]
                        log(f"    エラー発生表 行[{row_idx+1}]: {' | '.join(row_texts)}")
                except Exception as log_e:
                    log(f"    エラー発生表内容のログ出力にも失敗: {log_e}")

                # 要素レベル警告として上位へ通知
                notify_warning(file_name, f"表処理をスキップ: {e}")

                pass

    return state.found_parent_count

def process_single_file(input_path, output_path) -> tuple[bool, int, str]:
    """
    1つのWordファイルを処理してテキストファイルに出力
    
    Args:
        input_path (str): 入力Wordファイルのパス
        output_path (str): 出力テキストファイルのパス
    
    Returns:
        tuple: (成功したか, PARENTセクション数, エラーメッセージ)
    """
    try:
        # 標準出力をファイルにリダイレクト
        with open(output_path, 'w', encoding='utf-8') as f:
            original_stdout = sys.stdout
            sys.stdout = f
            
            try:
                parent_count = extract_marked_sections(input_path)
            finally:
                sys.stdout = original_stdout
        
        # ファイルの末尾の空行を削除
        with open(output_path, 'r', encoding='utf-8') as f:
            content = f.read()
        content = content.rstrip(LINE_BREAK) 
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return (True, parent_count, None)
    
    except Exception as e:
        error_msg = str(e)

        # ファイル単位失敗として上位へ通知
        notify_file_error(input_path, f"ファイル処理に失敗: {error_msg}")

        return (False, 0, error_msg)

def main() -> int:
    """
    複数のWordファイルを一括処理するメイン関数
    """
    global log_file
    
    # 処理開始時刻（全ファイル共通）
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    
    # ログディレクトリを作成
    log_dir = Path(LOG_DIR)
    log_dir.mkdir(parents=True, exist_ok=True)
    
    # ログファイルを開く
    log_path = log_dir / f"process_{timestamp}.log"
    log_file = open(log_path, 'w', encoding='utf-8')
    
    try:
        log("="*70)
        log("Word文書マーカー抽出処理開始")
        log(f"処理開始時刻: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        log("="*70)
        
        # 入力・出力ディレクトリのパス
        input_dir = Path(INPUT_DIR)
        output_dir = Path(OUTPUT_DIR)
        
        log(f"入力ディレクトリ: {input_dir}")
        log(f"出力ディレクトリ: {output_dir}")
        log(f"ログディレクトリ: {log_dir}")
        log("")
        
        # 入力ディレクトリの存在確認
        if not input_dir.exists():
            error_msg = f"入力ディレクトリが存在しません: {input_dir}"
            log(error_msg, also_print=True)
            log("ディレクトリを作成してWordファイルを配置してください。", also_print=True)

            # ★追加（致命的として通知）
            notify_fatal(error_msg)

            # 終了コードを返す
            return EXIT_ERROR
        
        # 出力ディレクトリを作成
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # 処理対象ファイルを取得（一時ファイル除外）
        all_files = list(input_dir.glob(FILE_PATTERN))
        # ~$で始まる一時ファイルを除外
        target_files = [f for f in all_files if not f.name.startswith('~$')]
        
        log(f"検出された全ファイル数: {len(all_files)}件")
        log(f"処理対象ファイル数: {len(target_files)}件 (一時ファイル除外後)")
        
        if not target_files:
            error_msg = f"処理対象ファイルが見つかりません: {input_dir / FILE_PATTERN}"
            log(error_msg, also_print=True)

            # 致命的として通知
            notify_fatal(error_msg)

            # 終了コードを返す
            return EXIT_ERROR
        
        print("="*70)
    
        # 処理結果を集計
        success_count = 0
        error_count = 0
        total_parent_sections = 0
        
        log("")
        log("ファイル処理開始")
        log("="*70)
        
        # 各ファイルを処理
        for idx, input_file in enumerate(target_files, 1):
            log(f"")
            log(f"[{idx}/{len(target_files)}] 処理中: {input_file.name}")
            print(f"[{idx}/{len(target_files)}] 処理中: {input_file.name}")
            
            # 仮の出力ファイル名を生成（エラー時に変更される可能性あり）
            output_filename = f"{input_file.stem}_{timestamp}.txt"
            output_path = output_dir / output_filename
            
            log(f"  入力ファイル: {input_file}")
            log(f"  出力ファイル: {output_path}")
            
            # 各ファイル処理
            success, parent_count, error_msg = process_single_file(str(input_file), str(output_path))
            
            if success:
                log(f"  結果: 成功")
                log(f"  [PARENT]セクション数: {parent_count}件")
                print(f"  [OK] 成功: {output_filename} ([PARENT]セクション: {parent_count}件)")
                success_count += 1
                total_parent_sections += parent_count
            else:
                # エラー時はファイル名を変更
                error_filename = f"{input_file.stem}_{timestamp}_ERROR.txt"
                error_path = output_dir / error_filename
                
                # 既に作成されているファイルがあればリネーム
                if output_path.exists():
                    output_path.rename(error_path)
                    log(f"  ファイルをリネーム: {output_filename} -> {error_filename}")
                
                log(f"  結果: エラー")
                log(f"  エラー内容: {error_msg}")
                log(f"  エラーファイル: {error_filename}")
                
                print(f"     エラー: {error_msg}")
                print(f"     エラーファイル: {error_filename}")
                error_count += 1
        
        # 最終結果サマリー
        log("")
        log("="*70)
        log("処理完了")
        log(f"処理終了時刻: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        log(f"成功: {success_count}件")
        log(f"エラー: {error_count}件")
        log(f"総[PARENT]セクション数: {total_parent_sections}件")
        log(f"出力先: {output_dir.resolve()}")
        log(f"ログファイル: {log_path.resolve()}")
        log("="*70)
        
        print("="*70)
        print(f"{LINE_BREAK}処理完了")
        print(f"  成功: {success_count}件")
        print(f"  エラー: {error_count}件")
        print(f"  総[PARENT]セクション数: {total_parent_sections}件")
        print(f"{LINE_BREAK}出力先: {output_dir.resolve()}")
        print(f"ログファイル: {log_path.resolve()}")
    
        # 完走後の終了コード集約
        if had_warning or had_file_error or error_count > 0:
            return EXIT_WARNING
        return EXIT_OK

    finally:
        # ログファイルをクローズ
        if log_file:
            log_file.close()

# --- メイン処理 ---

if __name__ == "__main__":
    try:
        sys.exit(main())
    except Exception as e:
        msg = f"致命的エラーが発生しました: {e}"
        notify_fatal(msg)
        print(f"{LINE_BREAK}--- スタックトレース ---", file=sys.stderr)
        traceback.print_exc(file=sys.stderr)
        sys.exit(EXIT_ERROR)